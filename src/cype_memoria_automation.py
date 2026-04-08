#!/usr/bin/env python3
# Script para automatizar la inserción de tablas e imágenes desde fuentes DOCX/XLSX a una plantilla DOCX

from __future__ import annotations # Para permitir anotaciones de tipo más flexibles (Python 3.7+)

import argparse # Para parsear argumentos de línea de comando (ej: fuentes, plantilla, mapeo, salida)
import copy # Para copiar elementos XML de tablas DOCX sin modificar originales
from io import BytesIO # Para manejar imágenes en memoria sin archivos temporales (me guarda la imagen como bytes para no tener que escribirla a disco)
import re # Para trabajar con expresiones regulares en reglas de búsqueda (ej: disparadores, encabezados, nombres de archivo)
import tempfile # Para crear archivos temporales de imagen al convertir rangos Excel a PNG
from dataclasses import dataclass # Para definir la estructura de las reglas de sección de forma clara y concisa
from pathlib import Path # Para manejar rutas de archivos de forma más cómoda y segura (en lugar de strings)
from typing import Any, Iterable, List, Optional, Sequence, TYPE_CHECKING # Para anotaciones de tipo más precisas y para evitar importaciones circulares o pesadas en tiempo de ejecución

import yaml # Para cargar las reglas de extracción/inserción desde un archivo YAML de mapeo (ej: qué buscar, dónde insertarlo, cómo coincidirlo)
from docx import Document # Para trabajar con documentos DOCX (abrir fuentes, plantilla, manipular contenido, guardar resultado)
from docx.document import Document as DocxDocumentType # Para anotaciones de tipo específicas de documentos DOCX (evita confusión con Workbook de Excel)
from docx.oxml import OxmlElement # Para manipular directamente el XML subyacente de DOCX (ej: insertar tablas e imágenes en posiciones específicas)
from docx.oxml.table import CT_Tbl # Para identificar elementos de tabla en el XML de DOCX al iterar bloques de contenido
from docx.oxml.text.paragraph import CT_P # Para identificar elementos de párrafo en el XML de DOCX al iterar bloques de contenido
from docx.shared import Inches # Para definir el tamaño de las imágenes insertadas en la plantilla (ej: ancho en pulgadas)
from docx.table import Table # Para trabajar con tablas en DOCX (leer contenido, copiar estructuras, etc.)
from docx.text.paragraph import Paragraph # Para trabajar con párrafos en DOCX (leer texto, agregar imágenes, etc.)
from docx.enum.text import WD_ALIGN_PARAGRAPH # Para alinear imágenes insertadas en el centro del párrafo.

if TYPE_CHECKING: # Para evitar importaciones pesadas de openpyxl y excel2img en tiempo de ejecución, solo se importan si se necesitan (ej: reglas que buscan contenido Excel)
    from openpyxl import Workbook # Para trabajar con libros de Excel (abrir fuentes XLSX, leer tablas y rangos, etc.)
    from openpyxl.worksheet.table import Table as ExcelTable # Para trabajar con tablas definidas en Excel (leer referencias, etc.)


# Define una regla de sección que indica cómo extraer contenido de fuentes y dónde insertarlo
@dataclass
class SectionRule:
    id: str  # Identificador único de la regla
    target_placeholder: str  # Marcador en la plantilla donde insertar contenido
    source_trigger_regex: Optional[str] = None  # Regex para encontrar párrafo disparador en DOCX
    table_header_regex: Optional[str] = None  # Regex para coincidir encabezado de tabla
    table_offset_after_trigger: int = 1  # Desplazamiento de tabla respecto al disparador
    image_offset_after_trigger: int = 1  # Desplazamiento de imagen respecto al disparador
    match_mode: str = "single"  # "single" (primera coincidencia) o "all" (todas)
    source_file_regex: Optional[str] = None  # Regex para filtrar archivos fuente por nombre
    excel_table_name: Optional[str] = None  # Nombre de tabla Excel a buscar
    excel_range: Optional[str] = None  # Rango Excel explícito (ej: Sheet1!A1:C10)
    excel_sheet_name: Optional[str] = None  # Nombre de hoja Excel
    docx_image: bool = False  # Si True, busca imágenes en DOCX en lugar de tablas
    image_width_inches: float = 6.5  # Ancho de imagen al insertar en pulgadas
    source_stop_regex: Optional[str] = None  # Regex para detener la recolección de tablas
    text_after_trigger: bool = False  # Si True, extrae párrafos de texto tras el trigger
    text_until_regex: Optional[str] = None  # Regex que indica dónde detener la captura de texto
    include_trigger: bool = True  # Si False, no incluye el párrafo disparador en el texto extraído


# Itera párrafos y tablas en el documento en orden secuencial (para detectar disparadores y luego contenido relacionado)
def iter_block_items(document: DocxDocumentType) -> Iterable[Paragraph | Table]:
    """Yield paragraphs and tables in document order."""
    body = document.element.body # Accede al cuerpo del documento para iterar sobre sus elementos XML (párrafos y tablas) en el orden en que aparecen en el documento, lo que es crucial para detectar disparadores y luego el contenido relacionado (tablas o imágenes) que viene después.
    for child in body.iterchildren(): # Itera sobre los elementos hijos directos del cuerpo del documento (que pueden ser párrafos o tablas)
        if isinstance(child, CT_P):
            yield Paragraph(child, document) # Si el elemento es un párrafo, lo envuelve en un objeto Paragraph de python-docx y lo devuelve
        elif isinstance(child, CT_Tbl):
            yield Table(child, document) # Si el elemento es una tabla, lo envuelve en un objeto Table de python-docx y lo devuelve


# Verifica si la regla busca contenido en Excel
def _is_excel_rule(rule: SectionRule) -> bool:
    return bool(rule.excel_table_name or rule.excel_range) # Si la regla tiene un nombre de tabla Excel o un rango Excel definido, se considera una regla de contenido Excel. Esto es importante para determinar qué tipo de búsqueda realizar (tablas/rangos en libros de Excel vs tablas/imágenes en documentos DOCX) y cómo procesar los resultados (ej: convertir rangos Excel a imágenes).


# Verifica si la regla extrae texto de párrafos DOCX
def _is_docx_text_rule(rule: SectionRule) -> bool:
    return rule.text_after_trigger


# Extrae párrafos de texto entre trigger y stop en un documento DOCX
def _paragraph_has_image(paragraph: Paragraph) -> bool:
    """Returns True if the paragraph contains an embedded image."""
    blips = paragraph._p.findall(
        ".//*[@{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed]"
    )
    return len(blips) > 0


def _make_bullet_paragraph_xml(text: str) -> Any:
    """Creates a new paragraph XML element with bullet point style (•) in body text."""
    from docx.oxml.ns import qn
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    # Style: Normal/body text
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), "Normal")
    pPr.append(pStyle)
    # Indentation to simulate bullet
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    pPr.append(ind)
    p.append(pPr)
    # Run with bullet character + text
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = "•  " + text
    r.append(t)
    p.append(r)
    return p


def find_matching_text_in_document(source_doc: DocxDocumentType, rule: SectionRule) -> List[Paragraph]:
    """Devuelve párrafos entre regex disparador y parada.
    Omite párrafos que contienen imágenes embebidas para evitar conflictos de ID de relación.
    """
    if not rule.source_trigger_regex:
        return []

    trigger_pattern = re.compile(rule.source_trigger_regex)
    stop_pattern = re.compile(rule.text_until_regex) if rule.text_until_regex else None

    matches: List[Paragraph] = []
    seen_trigger = False

    for block in iter_block_items(source_doc):
        if not isinstance(block, Paragraph):
            continue
        text = block.text.strip()

        if not seen_trigger:
            if text and trigger_pattern.search(text):
                seen_trigger = True
                if rule.include_trigger:
                    if not _paragraph_has_image(block):
                        matches.append(block)
            continue

        # Si hay stop y el texto coincide, termina
        if stop_pattern and text and stop_pattern.search(text):
            break

        # Ignorar párrafos con imágenes embebidas (evita conflicto de rId en la plantilla)
        if _paragraph_has_image(block):
            continue

        matches.append(block)

    return matches


# Selecciona párrafos de texto según modo de coincidencia
def select_docx_text_for_rule(
    source_docs: Sequence[tuple[Path, DocxDocumentType]], rule: SectionRule
) -> List[tuple[Paragraph, Path]]:
    selected: List[tuple[Paragraph, Path]] = []

    for source_path, source_doc in source_docs:
        if not _matches_source_file(source_path, rule.source_file_regex):
            continue
        matches = find_matching_text_in_document(source_doc, rule)
        for para in matches:
            selected.append((para, source_path))
        if selected and rule.match_mode == "single":
            return selected

    return selected


# Verifica si la regla busca imágenes DOCX
def _is_docx_image_rule(rule: SectionRule) -> bool:
    return rule.docx_image # Si la regla tiene docx_image=True, se considera una regla que busca imágenes en documentos DOCX en lugar de tablas. Esto afecta la lógica de búsqueda e inserción, ya que las imágenes se manejan de manera diferente a las tablas (ej: extraer imágenes incrustadas, insertarlas como imágenes en la plantilla, etc.).


# Importa dinámicamente openpyxl si no está disponible (lanza error)
def _require_openpyxl() -> Any: # Para evitar que la dependencia de openpyxl sea obligatoria si no se usan reglas de contenido Excel, se importa dinámicamente solo cuando se necesita. Si no está instalada, se lanza un error claro indicando que falta la dependencia y cómo instalarla.
    try:
        from openpyxl import load_workbook as _load_workbook
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "Falta la dependencia 'openpyxl'. Instalá dependencias con: pip install -r requirements.txt"
        ) from exc
    return _load_workbook # Devuelve la función load_workbook para que pueda ser usada en el resto del código sin importar openpyxl globalmente


# Carga las reglas desde un archivo YAML de mapeo
def load_rules(mapping_path: Path) -> List[SectionRule]: #Defino mis reglas como lista de objetos SectionRule, que se construyen a partir de la información cargada desde el archivo YAML. Cada sección en el YAML representa una regla que indica qué buscar en los documentos fuente y dónde insertarlo en la plantilla. Esta función se encarga de leer el YAML, validar su contenido y convertirlo en una lista de objetos SectionRule que luego serán usados para procesar los documentos.
    # Lee el YAML
    with mapping_path.open("r", encoding="utf-8") as fh:
        raw = yaml.safe_load(fh) or {}

    sections = raw.get("sections", [])
    rules: List[SectionRule] = []
    # Para cada sección en el YAML, crea una regla
    for i, section in enumerate(sections, start=1):
        source_trigger_regex = section.get("source_trigger_regex")
        table_header_regex = section.get("table_header_regex")
        excel_table_name = section.get("excel_table_name")
        excel_range = section.get("excel_range")
        docx_image = bool(section.get("docx_image", False))

        # Valida que al menos un criterio de búsqueda esté presente
        text_after_trigger = bool(section.get("text_after_trigger", False))
        if not excel_table_name and not excel_range and not source_trigger_regex and not table_header_regex and not docx_image and not text_after_trigger:
            raise ValueError(
                f"La sección {section.get('id', f'section_{i}')!r} debe tener al menos uno: "
                "'source_trigger_regex', 'table_header_regex', 'excel_table_name', 'excel_range' o 'docx_image'."
            )

        # Valida el modo de coincidencia
        match_mode = str(section.get("match_mode", "single")).lower().strip()
        if match_mode not in {"single", "all"}:
            raise ValueError(
                f"La sección {section.get('id', f'section_{i}')!r} tiene match_mode inválido: {match_mode!r}. "
                "Valores permitidos: 'single' o 'all'."
            )

        # Crea la regla con parámetros del YAML
        rules.append(
            SectionRule(
                id=section.get("id", f"section_{i}"),
                source_trigger_regex=source_trigger_regex,
                table_header_regex=table_header_regex,
                table_offset_after_trigger=int(section.get("table_offset_after_trigger", 1)),
                image_offset_after_trigger=int(section.get("image_offset_after_trigger", 1)),
                target_placeholder=section["target_placeholder"],
                match_mode=match_mode,
                source_file_regex=section.get("source_file_regex"),
                excel_table_name=excel_table_name,
                excel_range=excel_range,
                excel_sheet_name=section.get("excel_sheet_name"),
                docx_image=docx_image,
                image_width_inches=float(section.get("image_width_inches", 6.5)),
                source_stop_regex=section.get("source_stop_regex"),
                text_after_trigger=bool(section.get("text_after_trigger", False)),
                text_until_regex=section.get("text_until_regex"),
                include_trigger=bool(section.get("include_trigger", True)),
            )
        )
    return rules


# Verifica si el nombre del archivo coincide con la regex de filtro (si existe)
def _matches_source_file(path: Path, source_file_regex: Optional[str]) -> bool:
    if not source_file_regex:
        return True
    return bool(re.search(source_file_regex, path.name))


# Extrae el texto de la primera fila de una tabla (para identificarla)
def get_table_header_text(table: Table) -> str:
    """Return a normalized string from the first row of the table."""
    if not table.rows:
        return ""
    first_row = table.rows[0]
    values = [cell.text.strip() for cell in first_row.cells if cell.text.strip()]
    return " | ".join(values)


# Verifica si el encabezado de una tabla coincide con la regex especificada
def matches_table_header(table: Table, header_regex: Optional[str]) -> bool:
    if not header_regex:
        return True
    header_text = get_table_header_text(table)
    return bool(header_text and re.search(header_regex, header_text))


# Busca todas las tablas DOCX que coincidan con la regla en un documento
def find_matching_tables_in_document(source_doc: DocxDocumentType, rule: SectionRule) -> List[Table]:
    """Returns all matching DOCX tables in one document."""
    matches: List[Table] = []

    # Si hay disparador (trigger), busca párrafos que coincidan y luego tablas después
    if rule.source_trigger_regex:
        trigger_pattern = re.compile(rule.source_trigger_regex)
        stop_pattern = re.compile(rule.source_stop_regex) if getattr(rule, "source_stop_regex", None) else None
        seen_trigger = False

        for block in iter_block_items(source_doc):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                # Detecta el párrafo disparador
                if not seen_trigger and text and trigger_pattern.search(text):
                    seen_trigger = True
                    continue
                # Detecta el párrafo de parada
                if seen_trigger and stop_pattern and text and stop_pattern.search(text):
                    break

            # Tras el disparador, recolecta tablas que coincidan con encabezado
            if seen_trigger and isinstance(block, Table) and matches_table_header(block, rule.table_header_regex):
                matches.append(block)

        return matches

    # Si no hay disparador, busca directamente tablas con encabezado coincidente
    for block in iter_block_items(source_doc):
        if isinstance(block, Table) and matches_table_header(block, rule.table_header_regex):
            matches.append(block)

    return matches


# Extrae imágenes incrustadas en un párrafo DOCX
def _extract_images_from_paragraph(paragraph: Paragraph) -> List[BytesIO]:
    images: List[BytesIO] = []
    # Busca elementos blip (referencias a imágenes)
    for blip in paragraph._p.xpath(".//*[local-name()='blip']"):
        embed_rel_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if not embed_rel_id:
            continue
        image_part = paragraph.part.related_parts.get(embed_rel_id)
        if image_part is None:
            continue
        # Agrega los bytes de imagen a lista
        images.append(BytesIO(image_part.blob))
    return images


# Busca todas las imágenes DOCX que coincidan con la regla
def find_matching_images_in_document(source_doc: DocxDocumentType, rule: SectionRule) -> List[BytesIO]:
    """Returns all matching DOCX images in one document."""
    matches: List[BytesIO] = []

    # Si hay disparador, busca párrafo trigger e imágenes después
    if rule.source_trigger_regex:
        trigger_pattern = re.compile(rule.source_trigger_regex)
        seen_trigger = False

        for block in iter_block_items(source_doc):
            if not isinstance(block, Paragraph):
                continue

            text = block.text.strip()
            if text and trigger_pattern.search(text):
                seen_trigger = True
                continue

            # Tras el disparador, extrae imágenes
            if seen_trigger:
                matches.extend(_extract_images_from_paragraph(block))

        return matches

    # Sin disparador, extrae imágenes de todos los párrafos
    for block in iter_block_items(source_doc):
        if isinstance(block, Paragraph):
            matches.extend(_extract_images_from_paragraph(block))

    return matches


# Busca una tabla Excel por nombre en un libro de trabajo
def _find_excel_table_by_name(workbook: Workbook, table_name: str, sheet_name: Optional[str]) -> Optional[tuple[str, str]]:
    # Define hojas a buscar
    sheetnames = [sheet_name] if sheet_name else workbook.sheetnames
    for name in sheetnames:
        if name not in workbook.sheetnames:
            continue
        ws = workbook[name]
        # Busca tabla con nombre coincidente
        for candidate_name, candidate_table in ws.tables.items():
            if candidate_name == table_name or candidate_name.lower() == table_name.lower():
                table_ref = candidate_table.ref if hasattr(candidate_table, "ref") else str(candidate_table)
                return name, table_ref
    return None


# Resuelve un rango Excel desde la regla (ej: "Sheet1!A1:C10")
def _resolve_excel_range(rule: SectionRule) -> tuple[Optional[str], Optional[str]]:
    if not rule.excel_range:
        return None, None

    # Si el rango incluye "!", separa hoja y rango
    if "!" in rule.excel_range:
        sheet_name, range_ref = rule.excel_range.split("!", 1)
        return sheet_name.strip("'\""), range_ref

    return rule.excel_sheet_name, rule.excel_range


# Encuentra rangos Excel repetidos (múltiples bloques de datos con misma estructura)
def _find_repeated_excel_ranges(sheet: Any, template_range_ref: str) -> List[str]:
    from openpyxl.utils.cell import get_column_letter, range_boundaries

    # Calcula dimensiones del rango plantilla
    min_col, min_row, max_col, max_row = range_boundaries(template_range_ref)
    block_height = max_row - min_row + 1

    ranges: List[str] = []
    current_start_row = min_row

    # Itera hacia abajo encontrando bloques de datos
    while True:
        first_row_values = [
            sheet.cell(row=current_start_row, column=col_idx).value
            for col_idx in range(min_col, max_col + 1)
        ]
        # Si la fila está vacía, termina búsqueda
        if all(value is None or str(value).strip() == "" for value in first_row_values):
            break

        # Genera referencia de rango para este bloque
        current_end_row = current_start_row + block_height - 1
        range_ref = (
            f"{get_column_letter(min_col)}{current_start_row}:"
            f"{get_column_letter(max_col)}{current_end_row}"
        )
        ranges.append(range_ref)
        current_start_row += block_height

    return ranges


# Busca todos los rangos Excel que coincidan con la regla en un libro
def find_matching_excel_ranges_in_workbook(workbook: Workbook, rule: SectionRule) -> List[tuple[str, str]]:
    """Returns list of (sheet_name, range_ref) matching a rule in one workbook."""
    matches: List[tuple[str, str]] = []

    # Si hay nombre de tabla, búscala
    if rule.excel_table_name:
        found = _find_excel_table_by_name(workbook, rule.excel_table_name, rule.excel_sheet_name)
        if found:
            sheet_name, table_ref = found
            matches.append((sheet_name, table_ref))

    # Si hay rango, búscalo (posiblemente múltiples instancias)
    if rule.excel_range:
        sheet_name, range_ref = _resolve_excel_range(rule)
        if sheet_name and range_ref and sheet_name in workbook.sheetnames:
            for detected_range in _find_repeated_excel_ranges(workbook[sheet_name], range_ref):
                matches.append((sheet_name, detected_range))

    return matches


# Importa dinámicamente excel2img
def _require_excel2img() -> Any:
    try:
        import excel2img
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "Falta la dependencia 'excel2img'. Instalá dependencias con: pip install -r requirements.txt"
        ) from exc
    return excel2img


# Convierte un rango Excel a imagen PNG
def render_excel_range_as_image(source_xlsx_path: Path, sheet_name: str, range_ref: str) -> BytesIO:
    excel2img = _require_excel2img()

    # Crea archivo temporal para la imagen
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
        tmp_path = Path(tmp.name)

    try:
        # Exporta el rango como PNG
        excel2img.export_img(str(source_xlsx_path), str(tmp_path), sheet_name, range_ref)
        image_bytes = tmp_path.read_bytes()
    except Exception as exc:
        raise RuntimeError(
            f"No se pudo exportar como imagen el rango {sheet_name}!{range_ref} de {source_xlsx_path}. "
            "Verificá que Excel esté instalado y que el archivo no esté bloqueado."
        ) from exc
    finally:
        # Elimina archivo temporal
        if tmp_path.exists():
            tmp_path.unlink()

    return BytesIO(image_bytes)


# Selecciona tablas DOCX según modo de coincidencia (single o all)
def select_docx_tables_for_rule(
    source_docs: Sequence[tuple[Path, DocxDocumentType]], rule: SectionRule
) -> List[tuple[Table, Path]]:
    selected: List[tuple[Table, Path]] = []

    # Modo "single": toma primera coincidencia en primer archivo que la tenga
    if rule.match_mode == "single":
        for source_path, source_doc in source_docs:
            if not _matches_source_file(source_path, rule.source_file_regex):
                continue
            matches = find_matching_tables_in_document(source_doc, rule)
            idx = max(rule.table_offset_after_trigger - 1, 0)
            if idx < len(matches):
                selected.append((matches[idx], source_path))
                return selected
        return selected

    # Modo "all": toma todas las coincidencias de todos los archivos
    idx = max(rule.table_offset_after_trigger - 1, 0)
    for source_path, source_doc in source_docs:
        if not _matches_source_file(source_path, rule.source_file_regex):
            continue
        matches = find_matching_tables_in_document(source_doc, rule)
        if idx < len(matches):
            for table in matches[idx:]:
                selected.append((table, source_path))

    return selected


# Selecciona imágenes DOCX según modo de coincidencia
def select_docx_images_for_rule(
    source_docs: Sequence[tuple[Path, DocxDocumentType]], rule: SectionRule
) -> List[tuple[BytesIO, Path]]:
    selected: List[tuple[BytesIO, Path]] = []

    # Modo "single": primera imagen de primer archivo que la tenga
    if rule.match_mode == "single":
        for source_path, source_doc in source_docs:
            if not _matches_source_file(source_path, rule.source_file_regex):
                continue
            matches = find_matching_images_in_document(source_doc, rule)
            idx = max(rule.image_offset_after_trigger - 1, 0)
            if idx < len(matches):
                selected.append((matches[idx], source_path))
                return selected
        return selected

    # Modo "all": todas las imágenes de todos los archivos
    idx = max(rule.image_offset_after_trigger - 1, 0)
    for source_path, source_doc in source_docs:
        if not _matches_source_file(source_path, rule.source_file_regex):
            continue
        matches = find_matching_images_in_document(source_doc, rule)
        if idx < len(matches):
            for image_stream in matches[idx:]:
                selected.append((image_stream, source_path))

    return selected


# Selecciona rangos Excel según modo de coincidencia
def select_excel_ranges_for_rule(
    workbooks: Sequence[tuple[Path, Workbook]], rule: SectionRule
) -> List[tuple[Path, str, str]]:
    selected: List[tuple[Path, str, str]] = []

    # Modo "single": primer rango del primer libro que lo tenga
    if rule.match_mode == "single":
        for source_path, workbook in workbooks:
            if not _matches_source_file(source_path, rule.source_file_regex):
                continue
            matches = find_matching_excel_ranges_in_workbook(workbook, rule)
            idx = max(rule.table_offset_after_trigger - 1, 0)
            if idx < len(matches):
                sheet_name, range_ref = matches[idx]
                selected.append((source_path, sheet_name, range_ref))
                return selected
        return selected

    # Modo "all": todos los rangos de todos los libros
    idx = max(rule.table_offset_after_trigger - 1, 0)
    for source_path, workbook in workbooks:
        if not _matches_source_file(source_path, rule.source_file_regex):
            continue
        matches = find_matching_excel_ranges_in_workbook(workbook, rule)
        if idx < len(matches):
            for sheet_name, range_ref in matches[idx:]:
                selected.append((source_path, sheet_name, range_ref))

    return selected


# Busca el párrafo que contiene el marcador en la plantilla
def find_placeholder_paragraph(document: DocxDocumentType, placeholder: str) -> Optional[Paragraph]:
    for paragraph in document.paragraphs:
        if placeholder in paragraph.text:
            return paragraph
    return None


# Inserta tablas e imágenes en el lugar del marcador
def replace_placeholder_with_content(
    template_doc: DocxDocumentType,
    placeholder: str,
    docx_tables_to_insert: List[Table],
    docx_images_to_insert: List[BytesIO],
    excel_images_to_insert: List[BytesIO],
    image_width_inches: float,
    docx_paragraphs_to_insert: Optional[List[Paragraph]] = None,
) -> bool:
    # Busca el párrafo con el marcador
    paragraph = find_placeholder_paragraph(template_doc, placeholder)
    if paragraph is None:
        return False

    # Obtiene posición del párrafo en el documento
    parent = paragraph._p.getparent()
    paragraph_index = parent.index(paragraph._p)

    insertion_index = paragraph_index + 1
    paras_to_insert = docx_paragraphs_to_insert or []
    total_items = len(docx_tables_to_insert) + len(docx_images_to_insert) + len(excel_images_to_insert)
    current_item = 0

    # Inserta párrafos de texto como bullets de punto en estilo cuerpo
    for para in paras_to_insert:
        text = para.text.strip()
        if not text:
            continue
        bullet_xml = _make_bullet_paragraph_xml(text)
        parent.insert(insertion_index, bullet_xml)
        insertion_index += 1

    # Inserta las tablas DOCX
    for table in docx_tables_to_insert:
        current_item += 1
        new_table_xml = copy.deepcopy(table._tbl)
        parent.insert(insertion_index, new_table_xml)
        insertion_index += 1

        # Agrega espaciador entre items si no es el último
        if current_item < total_items:
            spacer_paragraph = OxmlElement("w:p")
            parent.insert(insertion_index, spacer_paragraph)
            insertion_index += 1

    # Inserta las imágenes (DOCX y Excel)
    for image_stream in [*docx_images_to_insert, *excel_images_to_insert]:
        current_item += 1

        # Crea párrafo para la imagen
        image_paragraph_xml = OxmlElement("w:p")
        parent.insert(insertion_index, image_paragraph_xml)
        image_paragraph = Paragraph(image_paragraph_xml, template_doc)
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Agrega imagen centrada
        run = image_paragraph.add_run()
        image_stream.seek(0)
        run.add_picture(image_stream, width=Inches(image_width_inches))
        insertion_index += 1

        # Espaciador
        if current_item < total_items:
            spacer_paragraph = OxmlElement("w:p")
            parent.insert(insertion_index, spacer_paragraph)
            insertion_index += 1

    # Limpia el marcador del párrafo (o lo elimina si queda vacío)
    paragraph.text = paragraph.text.replace(placeholder, "").strip()
    if not paragraph.text:
        parent.remove(paragraph._p)

    return True


# Recopila todos los archivos fuente desde argumentos
def list_source_files(
    source_docx: Optional[List[Path]], source_xlsx: Optional[List[Path]], source_dir: Optional[Path]
) -> List[Path]:
    source_files: List[Path] = []

    # Agrega archivos directamente especificados
    if source_docx:
        source_files.extend(source_docx)
    if source_xlsx:
        source_files.extend(source_xlsx)

    # Agrega archivos desde carpeta
    if source_dir:
        if not source_dir.exists() or not source_dir.is_dir():
            raise ValueError(f"La carpeta de fuentes no existe o no es válida: {source_dir}")
        source_files.extend(sorted(source_dir.glob("*.docx")))
        source_files.extend(sorted(source_dir.glob("*.xlsx")))

    # Elimina duplicados manteniendo orden
    unique: List[Path] = []
    seen = set()
    for path in source_files:
        resolved = str(path.resolve())
        if resolved not in seen:
            seen.add(resolved)
            unique.append(path)

    # Validaciones
    if not unique:
        raise ValueError("No se encontraron fuentes. Usá --source-dir, --source-docx y/o --source-xlsx.")

    missing = [str(p) for p in unique if not p.exists()]
    if missing:
        raise ValueError(f"No existen estos archivos fuente: {', '.join(missing)}")

    invalid_ext = [str(p) for p in unique if p.suffix.lower() not in {".docx", ".xlsx"}]
    if invalid_ext:
        raise ValueError(f"Extensiones no soportadas (solo .docx/.xlsx): {', '.join(invalid_ext)}")

    return unique


# Función principal que orquesta todo el proceso
def run(
    source_docx: Optional[List[Path]],
    source_xlsx: Optional[List[Path]],
    source_dir: Optional[Path],
    template_docx: Path,
    mapping_yaml: Path,
    output_docx: Path,
) -> None:
    # 1. Carga reglas de mapeo
    rules = load_rules(mapping_yaml)
    if not rules:
        raise ValueError("No se encontraron secciones en el archivo de mapeo YAML.")

    # 2. Recopila archivos fuente
    source_files = list_source_files(source_docx, source_xlsx, source_dir)

    # 3. Abre todos los DOCX y XLSX fuente
    docx_sources = [(p, Document(str(p))) for p in source_files if p.suffix.lower() == ".docx"]
    load_workbook = _require_openpyxl()
    excel_sources = [(p, load_workbook(filename=str(p), data_only=True)) for p in source_files if p.suffix.lower() == ".xlsx"]

    # 4. Abre la plantilla
    template_doc = Document(str(template_docx))

    # 5. Reporta fuentes cargadas
    print(f"Fuentes DOCX detectadas: {len(docx_sources)}")
    for source_file, _ in docx_sources:
        print(f"  - {source_file}")

    print(f"Fuentes Excel detectadas: {len(excel_sources)}")
    for source_file, _ in excel_sources:
        print(f"  - {source_file}")

    print(f"Reglas cargadas: {len(rules)}")

    # 6. Procesa cada regla
    for rule in rules:
        # Si es regla Excel
        if _is_excel_rule(rule):
            # Busca rangos Excel que coincidan
            selected_excel = select_excel_ranges_for_rule(excel_sources, rule)
            if not selected_excel:
                print(
                    f"[WARN] {rule.id}: no se encontró tabla Excel "
                    f"(mode={rule.match_mode!r}, file_regex={rule.source_file_regex!r}, "
                    f"table_name={rule.excel_table_name!r}, range={rule.excel_range!r}, "
                    f"sheet={rule.excel_sheet_name!r}, offset={rule.table_offset_after_trigger})"
                )
                continue

            # Convierte rangos a imágenes
            excel_images = [
                render_excel_range_as_image(source_path, sheet_name, range_ref)
                for source_path, sheet_name, range_ref in selected_excel
            ]
            sources = sorted(set(str(source_path) for source_path, _, _ in selected_excel))

            # Inserta imágenes en plantilla
            inserted = replace_placeholder_with_content(
                template_doc,
                rule.target_placeholder,
                docx_tables_to_insert=[],
                docx_images_to_insert=[],
                excel_images_to_insert=excel_images,
                image_width_inches=rule.image_width_inches,
            )
            if not inserted:
                print(f"[WARN] {rule.id}: no se encontró marcador {rule.target_placeholder!r} en plantilla")
                continue

            print(
                f"[OK] {rule.id}: {len(excel_images)} imagen(es) Excel insertadas en {rule.target_placeholder} "
                f"(fuentes: {', '.join(sources)})"
            )
            continue

        # Si es regla de texto DOCX
        if _is_docx_text_rule(rule):
            selected_text = select_docx_text_for_rule(docx_sources, rule)
            if not selected_text:
                print(
                    f"[WARN] {rule.id}: no se encontraron párrafos de texto "
                    f"(trigger={rule.source_trigger_regex!r}, stop={rule.text_until_regex!r})"
                )
                continue

            paragraphs = [para for para, _ in selected_text]
            sources = sorted(set(str(source) for _, source in selected_text))

            inserted = replace_placeholder_with_content(
                template_doc,
                rule.target_placeholder,
                docx_tables_to_insert=[],
                docx_images_to_insert=[],
                excel_images_to_insert=[],
                image_width_inches=rule.image_width_inches,
                docx_paragraphs_to_insert=paragraphs,
            )
            if not inserted:
                print(f"[WARN] {rule.id}: no se encontró marcador {rule.target_placeholder!r} en plantilla")
                continue

            print(
                f"[OK] {rule.id}: {len(paragraphs)} párrafo(s) insertados en {rule.target_placeholder} "
                f"(fuentes: {', '.join(sources)})"
            )
            continue

        # Si es regla de imágenes DOCX
        if _is_docx_image_rule(rule):
            selected_images = select_docx_images_for_rule(docx_sources, rule)
            if not selected_images:
                print(
                    f"[WARN] {rule.id}: no se encontró imagen DOCX "
                    f"(mode={rule.match_mode!r}, file_regex={rule.source_file_regex!r}, "
                    f"trigger={rule.source_trigger_regex!r}, offset={rule.image_offset_after_trigger})"
                )
                continue

            images = [image_stream for image_stream, _ in selected_images]
            sources = sorted(set(str(source) for _, source in selected_images))

            # Inserta imágenes en plantilla
            inserted = replace_placeholder_with_content(
                template_doc,
                rule.target_placeholder,
                docx_tables_to_insert=[],
                docx_images_to_insert=images,
                excel_images_to_insert=[],
                image_width_inches=rule.image_width_inches,
            )
            if not inserted:
                print(f"[WARN] {rule.id}: no se encontró marcador {rule.target_placeholder!r} en plantilla")
                continue

            print(
                f"[OK] {rule.id}: {len(images)} imagen(es) DOCX insertadas en {rule.target_placeholder} "
                f"(fuentes: {', '.join(sources)})"
            )
            continue

        # Si es regla de tablas DOCX
        selected_docx = select_docx_tables_for_rule(docx_sources, rule)
        if not selected_docx:
            print(
                f"[WARN] {rule.id}: no se encontró tabla DOCX "
                f"(mode={rule.match_mode!r}, file_regex={rule.source_file_regex!r}, "
                f"trigger={rule.source_trigger_regex!r}, header={rule.table_header_regex!r}, "
                f"offset={rule.table_offset_after_trigger})"
            )
            continue

        tables = [table for table, _ in selected_docx]
        sources = sorted(set(str(source) for _, source in selected_docx))

        # Inserta tablas en plantilla
        inserted = replace_placeholder_with_content(
            template_doc,
            rule.target_placeholder,
            docx_tables_to_insert=tables,
            docx_images_to_insert=[],
            excel_images_to_insert=[],
            image_width_inches=rule.image_width_inches,
        )
        if not inserted:
            print(f"[WARN] {rule.id}: no se encontró marcador {rule.target_placeholder!r} en plantilla")
            continue

        print(
            f"[OK] {rule.id}: {len(tables)} tabla(s) DOCX insertadas en {rule.target_placeholder} "
            f"(fuentes: {', '.join(sources)})"
        )

    # 7. Guarda documento final en DOCX
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    template_doc.save(str(output_docx))
    print(f"Archivo generado: {output_docx}")


# Parsea argumentos de línea de comando
def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Automatiza el pegado de tablas desde fuentes DOCX/XLSX de CYPECAD a plantilla de memoria propia."
    )
    parser.add_argument(
        "--source-docx",
        nargs="+",
        type=Path,
        help="Uno o más DOCX exportados desde CYPECAD (ej: --source-docx a.docx b.docx)",
    )
    parser.add_argument(
        "--source-xlsx",
        nargs="+",
        type=Path,
        help="Uno o más XLSX con tablas/rangos a importar (ej: --source-xlsx planilla1.xlsx)",
    )
    parser.add_argument(
        "--source-dir",
        type=Path,
        help="Carpeta con múltiples fuentes; se procesan todos los *.docx y *.xlsx",
    )
    parser.add_argument("--template-docx", required=True, type=Path, help="Plantilla DOCX de memoria de cálculo")
    parser.add_argument("--mapping-yaml", required=True, type=Path, help="YAML con reglas de extracción/inserción")
    parser.add_argument("--output-docx", required=True, type=Path, help="Ruta del DOCX final")
    return parser.parse_args()


# Punto de entrada
def main() -> None:
    args = parse_args()
    run(
        source_docx=args.source_docx,
        source_xlsx=args.source_xlsx,
        source_dir=args.source_dir,
        template_docx=args.template_docx,
        mapping_yaml=args.mapping_yaml,
        output_docx=args.output_docx,
    )


if __name__ == "__main__":
    main()
